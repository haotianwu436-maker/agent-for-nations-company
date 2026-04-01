from pathlib import Path

from docx import Document
from fpdf import FPDF


EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(parents=True, exist_ok=True)


def export_markdown_to_docx(job_id: str, markdown: str) -> Path:
    doc = Document()
    for line in markdown.splitlines():
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", "", 1), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    out = EXPORT_DIR / f"report-{job_id}.docx"
    doc.save(out)
    return out


def export_markdown_to_pdf(job_id: str, markdown: str) -> Path:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    width = 190
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            pdf.ln(4)
            continue
        text = line.replace("## ", "")
        text = text.replace("\t", " ").strip()
        if not text:
            continue
        try:
            pdf.multi_cell(width, 6, text=text)
        except Exception:
            safe = text.encode("latin-1", "ignore").decode("latin-1")
            if safe.strip():
                pdf.multi_cell(width, 6, text=safe)
    out = EXPORT_DIR / f"report-{job_id}.pdf"
    pdf.output(str(out))
    return out
